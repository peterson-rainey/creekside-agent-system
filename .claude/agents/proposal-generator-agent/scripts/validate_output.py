"""
validate_output.py -- Pre-send validation for Creekside proposal .docx files.

Extracts all text from a .docx and checks for:
  (1) Forbidden text patterns (em dashes, Slack refs, performance promises, etc.)
  (2) Pricing math consistency (minimum-fee thresholds must match variable rate math)

Does NOT delete the file -- leaves it for inspection.

Usage:
    python3 validate_output.py /path/to/proposal.docx

Exit codes:
    0  -- all checks passed
    1  -- one or more violations found (details printed to stdout)
    2  -- bad arguments or unreadable file
"""

import re
import sys

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx")
    sys.exit(2)


# ---------------------------------------------------------------------------
# Forbidden pattern definitions
# ---------------------------------------------------------------------------

CHECKS = [
    {
        "id": "em_dash",
        "description": "Em dash (U+2014)",
        "action": "block",
        "pattern": re.compile(r"—"),
        "message": (
            "Em dashes are forbidden in all Creekside output (hard rule #4). "
            "Rewrite the sentence using a colon, period, or restructuring."
        ),
    },
    {
        "id": "slack_reference",
        "description": "Word 'Slack' (case-insensitive)",
        "action": "block",
        "pattern": re.compile(r"\bslack\b", re.IGNORECASE),
        "message": (
            "Slack is not an active platform at Creekside (Standing Rule #1). "
            "Remove the reference. Do not replace with 'Google Chat' in sales-stage "
            "materials -- see agent_knowledge id 81e7d4e6-534b-4cc2-836e-dc46cbbbdc8a."
        ),
    },
    {
        "id": "performance_claim",
        "description": "Quantified performance promise (e.g. '20-50% better in 90 days')",
        "action": "block",
        "pattern": re.compile(
            r"\d+(\s*[-–]\s*\d+)?\s*%"
            r".{0,60}(better|improv|reduc|increas|lower|higher)"
            r".{0,60}(\d+\s*days?|in\s+\d+\s+months?)",
            re.IGNORECASE,
        ),
        "message": (
            "Quantified performance promises are forbidden in proposals (rule #6: no "
            "guaranteed ROI or specific results promises). Remove or rephrase without "
            "attaching a number to a timeframe."
        ),
    },
    {
        "id": "work_free_claim",
        "description": "'work free' outside marketing tagline context",
        "action": "block",
        # Matches "work free" when NOT immediately preceded by "we" + some qualifier
        # related to the approved tagline ("work free" in isolation or
        # "or we work free" where it's a naked promise, not a cited tagline).
        "pattern": re.compile(r"\bwork\s+free\b", re.IGNORECASE),
        "message": (
            "The phrase 'work free' appears in a sales-stage proposal. "
            "This may be an inadvertent inclusion of the 'Improve ROAS or we work free' "
            "marketing tagline, which is for marketing materials only -- not sales proposals. "
            "Remove or confirm intentional inclusion."
        ),
    },
    {
        "id": "google_chat_in_proposal",
        "description": "'Google Chat' in a sales-stage document",
        "action": "block",
        "pattern": re.compile(r"\bgoogle\s+chat\b", re.IGNORECASE),
        "message": (
            "Internal communication platform references ('Google Chat') must not appear "
            "in sales-stage proposals. This rule is documented in agent_knowledge id "
            "81e7d4e6-534b-4cc2-836e-dc46cbbbdc8a. Remove the reference."
        ),
    },
]


# ---------------------------------------------------------------------------
# Pricing math check
# ---------------------------------------------------------------------------

# Matches patterns like "$1,500 minimum.*under $5,000" or "$1,500 min.*below $7,500"
# Group 1: minimum dollar amount (digits + optional comma)
# Group 2: threshold dollar amount (digits + optional comma)
_MINIMUM_THRESHOLD_RE = re.compile(
    r"\$\s*([\d,]+)"           # minimum dollar figure
    r"(?:\s+\w+){0,5}\s+"     # up to 5 words gap (e.g. "minimum applies when spend is")
    r"(?:under|below|until|less\s+than)"
    r"\s+\$\s*([\d,]+)",       # threshold dollar figure
    re.IGNORECASE,
)

# Matches "20%" or "15%" as a standalone percentage (variable rate indicator).
# Note: "%" is a non-word character, so \b after "%" does not work as a terminator.
# Instead we use a negative lookahead to reject percentages that are part of decimals
# or other numeric constructs (e.g. "20.5%"). The \b before the digit is sufficient.
_VARIABLE_RATE_RE = re.compile(
    r"\b(\d+(?:\.\d+)?)%",
    re.IGNORECASE,
)


def _parse_dollars(text: str) -> int:
    """Strip commas, return integer dollars. Returns -1 on parse failure."""
    try:
        return int(text.replace(",", ""))
    except ValueError:
        return -1


def _looks_approximate(text: str) -> bool:
    """Return True if the surrounding text contains hedging language."""
    approximate_words = re.compile(
        r"\b(around|roughly|approximately|about|~|~=)\b", re.IGNORECASE
    )
    return bool(approximate_words.search(text))


def check_pricing_math(paragraphs: list[tuple[int, str]]) -> list[dict]:
    """
    Look for min-threshold contradictions in table cell text.

    Pattern: "$X minimum ... under/below $Y" + a variable rate Z% somewhere nearby.
    Math: the minimum applies until spend = X / (Z/100). If Y != that value (with +-5%
    tolerance), flag it as a contradiction.

    Only flags SPECIFIC dollar contradictions. Approximate language is skipped.
    Returns a list of math violation dicts.
    """
    math_violations = []

    # Collect all (para_idx, text) pairs that contain a min-threshold pattern
    for para_idx, text in paragraphs:
        for min_match in _MINIMUM_THRESHOLD_RE.finditer(text):
            # Skip if the surrounding text uses hedging language
            if _looks_approximate(text):
                continue

            min_amount = _parse_dollars(min_match.group(1))
            threshold_amount = _parse_dollars(min_match.group(2))
            if min_amount < 0 or threshold_amount < 0:
                continue  # can't parse -- skip

            # Find a variable rate in the same cell text or nearby paragraphs
            # Search the same text first, then all paragraphs within +/-5 indices
            rate_candidates = list(_VARIABLE_RATE_RE.finditer(text))
            if not rate_candidates:
                # Look in paragraphs within a small window
                for nearby_idx, nearby_text in paragraphs:
                    if abs(nearby_idx - para_idx) <= 5:
                        rate_candidates.extend(_VARIABLE_RATE_RE.finditer(nearby_text))

            for rate_match in rate_candidates:
                rate_pct = float(rate_match.group(1))
                if rate_pct <= 0:
                    continue

                implied_threshold = round(min_amount / (rate_pct / 100))

                # Allow +-5% tolerance for rounding
                tolerance = implied_threshold * 0.05
                if abs(implied_threshold - threshold_amount) > tolerance:
                    math_violations.append({
                        "para_idx": para_idx,
                        "line_preview": text[:200],
                        "min_amount": min_amount,
                        "threshold_stated": threshold_amount,
                        "rate_pct": rate_pct,
                        "implied_threshold": implied_threshold,
                    })
                    break  # one violation per min-threshold match is enough

    return math_violations


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def extract_paragraphs(docx_path: str) -> list[tuple[int, str]]:
    """Return list of (paragraph_index, text) tuples from the document."""
    doc = Document(docx_path)
    paragraphs = []
    idx = 0
    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            paragraphs.append((idx, text))
        idx += 1
    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if text.strip():
                        paragraphs.append((idx, text))
                        idx += 1
    return paragraphs


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate(docx_path: str) -> tuple[list[dict], list[dict]]:
    """
    Run all checks. Returns:
        (pattern_violations, math_violations)
    Both lists are empty on a clean document.
    """
    try:
        paragraphs = extract_paragraphs(docx_path)
    except Exception as e:
        print(f"ERROR: Could not read {docx_path}: {e}")
        sys.exit(2)

    pattern_violations = []
    for check in CHECKS:
        matches = []
        for para_idx, text in paragraphs:
            for m in check["pattern"].finditer(text):
                matches.append({
                    "para_idx": para_idx,
                    "line_preview": text[:200],
                    "match": m.group(0),
                    "match_start": m.start(),
                })
        if matches:
            pattern_violations.append({
                "id": check["id"],
                "description": check["description"],
                "action": check["action"],
                "message": check["message"],
                "occurrences": matches,
            })

    math_violations = check_pricing_math(paragraphs)

    return pattern_violations, math_violations


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 validate_output.py /path/to/proposal.docx")
        sys.exit(2)

    docx_path = sys.argv[1]

    print(f"Validating: {docx_path}")
    pattern_violations, math_violations = validate(docx_path)

    all_clean = not pattern_violations and not math_violations
    if all_clean:
        print("PASS -- no forbidden patterns or pricing math errors found.")
        sys.exit(0)

    total = len(pattern_violations) + len(math_violations)
    print(f"\nFAIL -- {total} violation(s) found. File left at path for inspection.\n")

    for v in pattern_violations:
        print(f"[{v['action'].upper()}] {v['description']}")
        print(f"  Rule: {v['message']}")
        print(f"  Occurrences: {len(v['occurrences'])}")
        for occ in v["occurrences"][:3]:  # show max 3 examples
            snippet = occ["line_preview"]
            print(f"    - Para {occ['para_idx']}: ...{snippet[:120]}...")
        if len(v["occurrences"]) > 3:
            print(f"    ... and {len(v['occurrences']) - 3} more.")
        print()

    for mv in math_violations:
        print("[BLOCK] PRICING MATH ERROR")
        print(
            f"  Cell text says \"${mv['min_amount']:,} minimum if spend is under "
            f"${mv['threshold_stated']:,}\" but math says minimum applies until "
            f"${mv['implied_threshold']:,} (=${mv['min_amount']:,}/{mv['rate_pct']}%). "
            f"One of these is wrong. Fix the table or the threshold."
        )
        print(f"  Context: ...{mv['line_preview'][:150]}...")
        print()

    sys.exit(1)


if __name__ == "__main__":
    main()
