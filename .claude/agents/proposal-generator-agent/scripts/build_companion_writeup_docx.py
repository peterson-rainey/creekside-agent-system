"""
build_companion_writeup_docx.py -- Creekside Marketing companion writeup builder.

Generates a branded standalone .docx companion document (e.g., conversion tracking
writeup, account audit findings, 90-day plan detail) that can be shared with a
prospect's partner, spouse, brother, team, or other stakeholder who wasn't on
the discovery call.

Usage:
    python3 build_companion_writeup_docx.py --json '{"lead_name": "...", ...}'
    python3 build_companion_writeup_docx.py  # reads JSON from stdin if no --json arg

Required JSON fields:
    lead_name           str     e.g. "Shin Nagpal"
    client_name         str     e.g. "Village Repair"
    topic               str     e.g. "Conversion Tracking Setup"
    subtitle            str     e.g. "What We Build for Village Repair"
    sections            list    list of dicts, each with:
                                  heading     str   section title
                                  body_paragraphs  list of str   paragraph text blocks
                                  bullet_lists     list of dicts with optional keys:
                                                     items  list of str  plain bullets
                                                     bold_prefix_items  list of [prefix, text]
    signature           str     e.g. "Peterson Rainey"

Optional JSON fields:
    out_path            str     full output path (default: ~/Desktop/proposals/{slug}_{topic}_{date}.docx)
    date_label          str     override date shown in "Prepared for" block (default: today)
    closing_paragraph   str     custom closing paragraph before signature
                                (default: standard questions/reply paragraph)
"""

import argparse
import json
import os
import re
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Brand constants (match build_lead_docx.py)
# ---------------------------------------------------------------------------
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def _slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]", "_", text).strip("_")


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

class CompanionWriteupBuilder:
    def __init__(self, params: dict):
        self.p = params
        self.doc = Document()
        self._setup_margins()

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    def _set_font(self, run, size=11, bold=False, color=None, italic=False):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def heading1(self, text):
        p = self.doc.add_heading(text, level=1)
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        return p

    def heading3(self, text):
        p = self.doc.add_heading(text, level=3)
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11.5)
        p.runs[0].font.color.rgb = DARK
        p.runs[0].font.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        return p

    def body(self, text, bold=False, italic=False, space_after=6):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, bold=bold, italic=italic)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r1 = p.add_run(bold_prefix + " ")
            self._set_font(r1, bold=True)
            r2 = p.add_run(text)
            self._set_font(r2)
        else:
            run = p.add_run(text)
            self._set_font(run)
        p.paragraph_format.space_after = Pt(3)
        return p

    def build_title(self):
        p = self.p
        title = self.doc.add_heading(p["topic"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = "Calibri"
        title.runs[0].font.size = Pt(22)
        title.runs[0].font.color.rgb = BLUE
        title.paragraph_format.space_after = Pt(2)

        sub = self.doc.add_paragraph(p.get("subtitle", ""))
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.name = "Calibri"
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.color.rgb = GRAY
        sub.paragraph_format.space_after = Pt(18)

    def build_prepared_for(self):
        p = self.p
        date_label = p.get("date_label", date.today().strftime("%B %d, %Y"))
        prep = self.doc.add_table(rows=1, cols=1)
        prep.alignment = WD_TABLE_ALIGNMENT.CENTER
        prep_cell = prep.rows[0].cells[0]
        shade_cell(prep_cell, "F9FAFB")
        prep_cell.text = ""
        p1 = prep_cell.paragraphs[0]
        r1a = p1.add_run("Prepared for: ")
        self._set_font(r1a, size=10.5, bold=True)
        r1b = p1.add_run(f"{p['lead_name']}, {p['client_name']}")
        self._set_font(r1b, size=10.5)
        p2 = prep_cell.add_paragraph()
        r2a = p2.add_run("Date: ")
        self._set_font(r2a, size=10.5, bold=True)
        r2b = p2.add_run(date_label)
        self._set_font(r2b, size=10.5)
        self.doc.add_paragraph()

    def build_sections(self):
        for section in self.p.get("sections", []):
            heading = section.get("heading", "")
            if heading:
                self.heading1(heading)

            for para in section.get("body_paragraphs", []):
                self.body(para)

            bullet_lists = section.get("bullet_lists", [])
            for blist in bullet_lists:
                # Plain bullets
                for item in blist.get("items", []):
                    self.bullet(item)
                # Bold-prefix bullets: each entry is [prefix, text]
                for item in blist.get("bold_prefix_items", []):
                    if isinstance(item, list) and len(item) == 2:
                        self.bullet(item[1], bold_prefix=item[0])
                    else:
                        self.bullet(str(item))

    def build_closing(self):
        p = self.p
        self.doc.add_paragraph()
        closing = p.get(
            "closing_paragraph",
            "If anything here is unclear or you want to dig deeper into how a specific "
            "piece works, just reply to the email and I'll walk through any of this in "
            "more detail."
        )
        self.body(closing)

    def build_signature(self):
        self.doc.add_paragraph()
        sig = self.doc.add_paragraph()
        r1 = sig.add_run(self.p.get("signature", "Peterson Rainey") + "\n")
        self._set_font(r1, size=11, bold=True)
        r2 = sig.add_run("Creekside Marketing")
        self._set_font(r2, size=11)

    def build(self) -> str:
        p = self.p
        client_slug = _slugify(p["client_name"])
        topic_slug = _slugify(p["topic"])
        today = date.today().strftime("%Y-%m-%d")

        out_path = p.get("out_path") or os.path.expanduser(
            f"~/Desktop/proposals/{client_slug}_{topic_slug}_{today}.docx"
        )
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        self.build_title()
        self.build_prepared_for()
        self.build_sections()
        self.build_closing()
        self.build_signature()

        self.doc.save(out_path)
        return out_path


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Build a Creekside companion writeup .docx"
    )
    parser.add_argument(
        "--json", dest="json_str",
        help="JSON string with writeup parameters (use stdin if omitted)"
    )
    args = parser.parse_args()

    if args.json_str:
        params = json.loads(args.json_str)
    else:
        params = json.load(sys.stdin)

    required = ["lead_name", "client_name", "topic", "subtitle", "sections", "signature"]
    missing = [f for f in required if f not in params]
    if missing:
        print(f"ERROR: Missing required fields: {missing}", file=sys.stderr)
        sys.exit(1)

    builder = CompanionWriteupBuilder(params)
    out_path = builder.build()
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
