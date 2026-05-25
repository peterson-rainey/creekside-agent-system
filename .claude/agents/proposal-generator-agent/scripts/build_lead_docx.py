"""
build_lead_docx.py -- Creekside Marketing per-lead proposal builder.

Generates a branded .docx proposal for a specific lead. This is a parameterized
version of build_shin_proposal_docx.py (the Village Repair reference build). The
master template builder is build_docx.py -- do not modify that script here.

Usage:
    python3 build_lead_docx.py --json '{"lead_name": "...", ...}'
    python3 build_lead_docx.py  # reads JSON from stdin if no --json arg

Required JSON fields:
    lead_name           str     e.g. "Shin Nagpal"
    client_name         str     e.g. "Village Repair"
    industry_context    str     e.g. "automotive repair shop"
    starting_ad_spend   int     monthly ad spend in dollars
    audit_findings      list    strings -- specific call-out points from discovery call
    signature           str     e.g. "Peterson Rainey"

Optional JSON fields:
    pricing_override    dict    override pricing figures (keys: onboarding_fee,
                                monthly_min, variable_rate_desc, monthly_cap)
    out_path            str     full output path (default: ~/Desktop/proposals/{slug}_{date}.docx)
    proposal_title      str     override document title (default: "Google Ads Management Proposal")
    audit_findings_section  dict    if present, renders a shaded "What We Found in Your Account"
                                callout at the top of the Onboarding section. Keys:
                                  findings (list of str -- specific account observations)
                                  header   (str, optional -- defaults to "What We Found in Your Account")
                                Use this for Step 2.5 audit-finding injection. Distinct from the
                                legacy audit_findings list (which appears in Overview body text).
    sections            dict    override section body text (keys: overview, why_choose, next_steps)
    onboarding_items    list    dicts with keys: heading (str) and bullets (list of str or
                                [bold_prefix, text] pairs)
    ongoing_items       list    additional ongoing management bullet dicts
    payment_terms       str     override payment/contract terms description
    date_label          str     override date shown in "Prepared for" block (default: today)
"""

import argparse
import json
import os
import re
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Current pricing (2026-05-25) -- matches pricing-logic.md and build_docx.py
# Simplified from 3-plan structure: Plan B Shared and Plan C Retainer removed.
DEFAULT_PRICING = {
    "onboarding_fee": "$1,500 per platform (one-time)",
    "monthly_min": "$1,500 minimum per platform",
    "variable_rate_desc": (
        "20% of ad spend up to $30,000/month\n"
        "15% from $30,000 to $60,000\n"
        "10% above $60,000\n"
        "(Calculated per platform; $1,500 minimum applies until spend exceeds $7,500/mo)"
    ),
    "monthly_cap": "$15,000 / month",
    "plan_label": "Creekside Management Fee",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]", "_", text).strip("_")


def _derive_plan_mode(starting_ad_spend: int, recommended_plan: str) -> str:
    """Return 'single', 'two', or 'all' based on spend thresholds."""
    if starting_ad_spend < SINGLE_PLAN_THRESHOLD:
        return "single"
    if starting_ad_spend < TWO_PLAN_THRESHOLD:
        return "two"
    return "all"


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Document builder class
# ---------------------------------------------------------------------------

class ProposalBuilder:
    def __init__(self, params: dict):
        self.p = params
        self.doc = Document()
        self._setup_margins()

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    # --- Low-level formatters ---

    def _set_font(self, run, size=11, bold=False, color=None, italic=False):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def heading1(self, text):
        p = self.doc.add_heading(text, level=1)
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        return p

    def heading3(self, text):
        p = self.doc.add_heading(text, level=3)
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11.5)
        p.runs[0].font.color.rgb = DARK
        p.runs[0].font.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        return p

    def body(self, text, bold=False, italic=False, space_after=6):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, bold=bold, italic=italic)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r1 = p.add_run(bold_prefix + " ")
            self._set_font(r1, bold=True)
            r2 = p.add_run(text)
            self._set_font(r2)
        else:
            run = p.add_run(text)
            self._set_font(run)
        p.paragraph_format.space_after = Pt(3)
        return p

    def numbered(self, text):
        p = self.doc.add_paragraph(style="List Number")
        run = p.add_run(text)
        self._set_font(run)
        p.paragraph_format.space_after = Pt(3)
        return p

    def set_cell_text(self, cell, text, bold=False, size=10, color=None,
                      align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    # --- Sections ---

    def build_title(self):
        p = self.p
        title_text = p.get("proposal_title", "Google Ads Management Proposal")
        title = self.doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = "Calibri"
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.color.rgb = BLUE
        title.paragraph_format.space_after = Pt(4)

        sub = self.doc.add_paragraph("Creekside Marketing")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.name = "Calibri"
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.color.rgb = GRAY
        sub.paragraph_format.space_after = Pt(18)

    def build_prepared_for(self):
        p = self.p
        date_label = p.get("date_label", date.today().strftime("%B %d, %Y"))
        prep = self.doc.add_table(rows=1, cols=1)
        prep.alignment = WD_TABLE_ALIGNMENT.CENTER
        prep_cell = prep.rows[0].cells[0]
        shade_cell(prep_cell, "F9FAFB")
        prep_cell.text = ""
        p1 = prep_cell.paragraphs[0]
        r1a = p1.add_run("Prepared for: ")
        self._set_font(r1a, size=10.5, bold=True)
        r1b = p1.add_run(f"{p['lead_name']}, {p['client_name']}")
        self._set_font(r1b, size=10.5)
        p2 = prep_cell.add_paragraph()
        r2a = p2.add_run("Date: ")
        self._set_font(r2a, size=10.5, bold=True)
        r2b = p2.add_run(date_label)
        self._set_font(r2b, size=10.5)
        self.doc.add_paragraph()

    def build_overview(self):
        p = self.p
        self.heading1("Overview")
        sections = p.get("sections", {})
        if "overview" in sections:
            self.body(sections["overview"])
        else:
            self.body(
                f"Thank you for the time recently. This proposal outlines how Creekside "
                f"Marketing will manage Google Ads for {p['client_name']}, with the goal of "
                f"driving qualified leads and revenue from paid advertising."
            )
            if p.get("audit_findings"):
                self.body(
                    "Based on our discovery conversation, we've noted the following key areas "
                    "to address:"
                )
                for finding in p["audit_findings"]:
                    self.bullet(finding)

    def build_why_choose(self):
        p = self.p
        self.heading1("Why Creekside Marketing")
        sections = p.get("sections", {})
        if "why_choose" in sections:
            self.body(sections["why_choose"])
        else:
            self.bullet(
                "You communicate directly with the people running your campaigns, not through "
                "a support inbox or account coordinator.",
                bold_prefix="Direct access to your ad managers."
            )
            self.bullet(
                "We optimize for revenue and booked appointments, not just cheap clicks or "
                "inflated impression counts.",
                bold_prefix="Growth-focused optimization."
            )
            self.bullet(
                "Google Ads and Meta Ads management. We don't do SEO, organic social, email "
                "marketing, or web design. If you need those, we'll tell you and recommend "
                "someone. We won't try to bolt them on.",
                bold_prefix="We do what we do."
            )

    def build_audit_findings_callout(self, findings: list, header: str = "What We Found in Your Account"):
        """
        Render a shaded callout block listing specific audit findings extracted from
        the discovery call. Only called when audit_findings_section is present and
        non-empty. Uses a single-cell table with a light-blue background to visually
        distinguish the block from body text.
        """
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        shade_cell(cell, "EBF5FB")  # light blue tint

        cell.text = ""
        p_header = cell.paragraphs[0]
        r_header = p_header.add_run(header)
        self._set_font(r_header, size=10.5, bold=True, color=BLUE)
        p_header.paragraph_format.space_after = Pt(4)

        for finding in findings:
            p_finding = cell.add_paragraph()
            r_bullet = p_finding.add_run("•  " + finding)
            self._set_font(r_bullet, size=10)
            p_finding.paragraph_format.space_after = Pt(2)
            p_finding.paragraph_format.left_indent = Inches(0.15)

        self.doc.add_paragraph()  # spacer after the callout

    def build_onboarding(self):
        p = self.p
        self.heading1("Onboarding Services")

        # Audit findings callout: rendered at the top of Onboarding when
        # audit_findings_section is explicitly provided (Step 2.5 of agent).
        # Falls back to the legacy audit_findings list in the Overview if not set.
        audit_section = p.get("audit_findings_section")
        if audit_section and isinstance(audit_section, dict):
            findings_list = audit_section.get("findings", [])
            section_header = audit_section.get("header", "What We Found in Your Account")
            if findings_list:
                self.build_audit_findings_callout(findings_list, header=section_header)

        onboarding_items = p.get("onboarding_items")
        if onboarding_items:
            for item in onboarding_items:
                self.heading3(item["heading"])
                for b in item.get("bullets", []):
                    if isinstance(b, list):
                        self.bullet(b[1], bold_prefix=b[0])
                    else:
                        self.bullet(b)
        else:
            # Default onboarding sections
            self.heading3("Conversion Tracking Setup")
            self.body(
                "This is often the most impactful work we do in month one. Without accurate "
                "conversion tracking, Google can't optimize for the outcomes that matter to "
                "your business. Specifically:",
                italic=True
            )
            self.bullet(
                "Track exactly which ads, keywords, and campaigns drive phone calls that "
                "convert to paying customers.",
                bold_prefix="Call tracking integration."
            )
            self.bullet(
                "Pull customer and appointment data back into Google Ads so the algorithm "
                "optimizes for people who actually convert -- not just anyone who clicks.",
                bold_prefix="CRM / booking system integration."
            )
            self.bullet(
                "Capture Google Click IDs through to revenue in your CRM, closing the "
                "attribution loop end-to-end.",
                bold_prefix="GCLID and UTM parameter capture."
            )
            self.bullet(
                "Enhanced accuracy and privacy compliance when required.",
                bold_prefix="Server-side conversion tracking."
            )

            self.heading3("Account Setup")
            self.bullet(
                f"Campaign architecture optimized for {p['industry_context']} in your "
                "target service area."
            )
            self.bullet("Keyword strategy built around high-intent search demand.")
            self.bullet("Ad copy written and tested for your audience.")

            self.heading3("90-Day Strategic Plan")
            self.bullet(
                "Conversion tracking live, campaigns launched, baseline data collected.",
                bold_prefix="Month 1: Foundation."
            )
            self.bullet(
                "A/B testing on ads, keyword refinement, budget reallocation toward "
                "what's converting.",
                bold_prefix="Month 2: Optimization."
            )
            self.bullet(
                "Scale what's working, add remarketing, evaluate budget expansion based "
                "on demand capture.",
                bold_prefix="Month 3: Performance and scale."
            )

    def build_ongoing(self):
        self.heading1("Ongoing Management")
        self.body("Once campaigns are live, we provide continuous hands-on management:")

        self.heading3("Reporting Cadence")
        self.bullet(
            "Comprehensive written reports with data analysis, insights, and strategic "
            "recommendations delivered every two weeks.",
            bold_prefix="Bi-weekly performance reports."
        )
        self.bullet(
            "Dedicated time to review performance, discuss results, and align on "
            "strategic priorities.",
            bold_prefix="Monthly strategy calls."
        )

        self.heading3("Optimization Activities")
        standard_bullets = [
            "Search term exclusions and negative keyword additions",
            "Daily review of Google Ads notifications for policy or account issues",
            "Ad placement analysis to exclude underperformers",
            "Disapproved ad and policy violation resolution",
            "Budget management and reallocation",
            "Conversion tracking verification (ongoing)",
            "Keyword bid and match type adjustments",
            "Cross-campaign performance analysis",
            "Ad group structure and targeting optimization",
        ]
        for b in standard_bullets:
            self.bullet(b)

        additional = self.p.get("ongoing_items", [])
        for item in additional:
            if isinstance(item, list):
                self.bullet(item[1], bold_prefix=item[0])
            else:
                self.bullet(item)

    def build_investment(self):
        p = self.p
        recommended = p["recommended_plan"].upper()
        spend = p.get("starting_ad_spend", 0)

        # Determine display mode
        if "single_plan_mode" in p and isinstance(p["single_plan_mode"], bool):
            explicit_mode = p["single_plan_mode"]
            if explicit_mode:
                mode = "single"
            else:
                # Use spend to determine two vs all
                mode = _derive_plan_mode(spend, recommended)
                if mode == "single":
                    mode = "two"  # override: caller said don't be single
        else:
            mode = _derive_plan_mode(spend, recommended)

        self.doc.add_page_break()
        self.heading1("Investment & Terms")

        pricing_override = p.get("pricing_override", {})

        if mode == "single":
            self._build_single_plan_table(recommended, pricing_override)
        elif mode == "two":
            self._build_two_plan_table(recommended, pricing_override)
        else:
            self._build_all_plans_table(pricing_override)

        self.doc.add_paragraph()
        self._build_payment_terms()

    def _get_pricing(self, plan_key: str, pricing_override: dict) -> dict:
        base = dict(DEFAULT_PRICING[plan_key])
        base.update(pricing_override)
        return base

    def _build_single_plan_table(self, plan: str, pricing_override: dict):
        pricing = self._get_pricing(plan, pricing_override)
        label = pricing["plan_label"]

        self.heading3("Pricing Structure")
        self.body(
            "Your pricing is structured to scale directly with your ad spend, so "
            "you're not paying for capacity you don't need yet."
        )

        ROWS = [
            ["Onboarding Fee (one-time)", pricing["onboarding_fee"]],
            ["Monthly Management Fee", pricing["monthly_min"]],
            ["Variable Rate", pricing["variable_rate_desc"]],
            ["Monthly Cap", pricing["monthly_cap"]],
        ]
        col_widths = [Inches(2.6), Inches(3.6)]
        tbl = self.doc.add_table(rows=1 + len(ROWS), cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = tbl.rows[0].cells
        shade_cell(hdr_cells[0], "1A56DB")
        shade_cell(hdr_cells[1], "1A56DB")
        self.set_cell_text(
            hdr_cells[0], label, bold=True, size=11,
            color=WHITE, align=WD_ALIGN_PARAGRAPH.LEFT
        )
        hdr_cells[0].merge(hdr_cells[1])

        ROW_COLORS = ["F9FAFB", "FFFFFF", "F9FAFB", "FFFFFF"]
        for r_idx, row_data in enumerate(ROWS):
            cells = tbl.rows[r_idx + 1].cells
            for c_idx, cell_text in enumerate(row_data):
                shade_cell(cells[c_idx], ROW_COLORS[r_idx % len(ROW_COLORS)])
                bold = (c_idx == 0)
                self.set_cell_text(
                    cells[c_idx], cell_text, bold=bold,
                    size=10.5, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT
                )

        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]

    def _build_two_plan_table(self, recommended: str, pricing_override: dict):
        plans = ["A", "B"]
        self._build_multi_plan_table(plans, recommended, pricing_override)

    def _build_all_plans_table(self, pricing_override: dict):
        self._build_multi_plan_table(["A", "B", "C"], None, pricing_override)

    def _build_multi_plan_table(self, plans: list, recommended: str, pricing_override: dict):
        self.heading3("Pricing Structure")
        self.body(
            "We offer multiple pricing plans. The recommended plan for "
            f"{self.p['client_name']} is marked below."
        )

        HEADER_ROW = ["Feature"] + [
            (DEFAULT_PRICING[pl]["plan_label"] +
             (" [Recommended]" if pl == recommended else ""))
            for pl in plans
        ]
        ROWS = [
            ["Onboarding Fee (one-time)"] + [self._get_pricing(pl, pricing_override)["onboarding_fee"] for pl in plans],
            ["Monthly Management Fee"] + [self._get_pricing(pl, pricing_override)["monthly_min"] for pl in plans],
            ["Variable Rate"] + [self._get_pricing(pl, pricing_override)["variable_rate_desc"] for pl in plans],
            ["Monthly Cap"] + [self._get_pricing(pl, pricing_override)["monthly_cap"] for pl in plans],
        ]

        num_cols = 1 + len(plans)
        col_widths_map = {
            2: [Inches(2.0), Inches(4.2)],
            3: [Inches(1.6), Inches(2.4), Inches(2.4)],
            4: [Inches(1.4), Inches(1.85), Inches(1.85), Inches(1.85)],
        }
        col_widths = col_widths_map.get(num_cols, [Inches(6.2 / num_cols)] * num_cols)

        tbl = self.doc.add_table(rows=1 + len(ROWS), cols=num_cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(HEADER_ROW):
            shade_cell(hdr_cells[i], "1A56DB")
            self.set_cell_text(
                hdr_cells[i], h, bold=True, size=10,
                color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER
            )

        ROW_COLORS = ["F9FAFB", "FFFFFF", "F9FAFB", "FFFFFF"]
        for r_idx, row_data in enumerate(ROWS):
            cells = tbl.rows[r_idx + 1].cells
            for c_idx, cell_text in enumerate(row_data):
                shade_cell(cells[c_idx], ROW_COLORS[r_idx % len(ROW_COLORS)])
                bold = (c_idx == 0)
                self.set_cell_text(
                    cells[c_idx], cell_text, bold=bold, size=9.5,
                    color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT
                )

        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]

    def _build_payment_terms(self):
        p = self.p
        self.heading3("Payment & Contract Terms")
        terms_override = p.get("payment_terms")
        if terms_override:
            self.body(terms_override)
        else:
            self.bullet(
                "You pay Google directly. Creekside never touches your ad budget.",
                bold_prefix="Ad spend:"
            )
            self.bullet(
                "Billed separately, paid directly to Creekside Marketing.",
                bold_prefix="Management fees:"
            )
            self.bullet(
                "3-month minimum contract. Month one is non-cancellable (full payment "
                "regardless of when you cancel). A $250 cancellation fee applies if you "
                "cancel during months 2 or 3. After the initial 3 months, service converts "
                "to month-to-month with 30-day cancellation notice.",
                bold_prefix="Contract terms:"
            )
            self.bullet(
                "All ad accounts, audiences, and creative assets remain your property. "
                "On cancellation, everything transfers back to you with full access.",
                bold_prefix="Account ownership:"
            )

    def build_next_steps(self):
        p = self.p
        self.heading1("Next Steps")
        sections = p.get("sections", {})
        if "next_steps" in sections:
            self.body(sections["next_steps"])
        else:
            self.numbered("Review this proposal and reach out with any questions.")
            self.numbered("Let us know you're ready to move forward.")
            self.numbered(
                f"We send over the onboarding invoice, service agreement, and onboarding "
                f"instructions (including what access we'll need from {p['client_name']})."
            )
            self.numbered("Kickoff call scheduled -- conversion tracking buildout begins.")
        self.body(
            "We're ready to start as soon as you give the green light. Reach out with "
            "any questions in the meantime."
        )

    def build_signature(self):
        self.doc.add_paragraph()
        sig = self.doc.add_paragraph()
        r1 = sig.add_run("Best regards,\n")
        self._set_font(r1, size=11)
        r2 = sig.add_run(self.p.get("signature", "Peterson Rainey") + "\n")
        self._set_font(r2, size=11, bold=True)
        r3 = sig.add_run("Creekside Marketing")
        self._set_font(r3, size=11)

    def build(self) -> str:
        p = self.p
        client_slug = _slugify(p["client_name"])
        today = date.today().strftime("%Y-%m-%d")

        out_path = p.get("out_path") or os.path.expanduser(
            f"~/Desktop/proposals/{client_slug}_{today}.docx"
        )
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        self.build_title()
        self.build_prepared_for()
        self.build_overview()
        self.build_why_choose()
        self.build_onboarding()
        self.build_ongoing()
        self.build_investment()
        self.build_next_steps()
        self.build_signature()

        self.doc.save(out_path)
        return out_path


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Build a Creekside per-lead proposal .docx")
    parser.add_argument(
        "--json", dest="json_str",
        help="JSON string with proposal parameters (use stdin if omitted)"
    )
    args = parser.parse_args()

    if args.json_str:
        params = json.loads(args.json_str)
    else:
        params = json.load(sys.stdin)

    # Validate required fields
    required = ["lead_name", "client_name", "industry_context",
                "starting_ad_spend", "recommended_plan", "signature"]
    missing = [f for f in required if f not in params]
    if missing:
        print(f"ERROR: Missing required fields: {missing}", file=sys.stderr)
        sys.exit(1)

    # Default audit_findings to empty list if not provided
    if "audit_findings" not in params:
        params["audit_findings"] = []

    builder = ProposalBuilder(params)
    out_path = builder.build()
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
