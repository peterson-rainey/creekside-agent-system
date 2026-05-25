"""
Build the Creekside Marketing Pricing Proposal as a .docx file.
Mirrors the live Google Doc template; pricing matches proposal_chart.py.

Run order: proposal_chart.py first (generates chart), then this script
(embeds the chart in the Investment & Terms section).

Simplified to single pricing plan on 2026-05-25 (Plan B Shared and Plan C Retainer removed).
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CHART_PATH = os.path.join(SCRIPT_DIR, "out", "proposal_chart.png")
OUT_PATH = os.path.join(SCRIPT_DIR, "out", "Pricing_Proposal_Creekside.docx")
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    p.runs[0].font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(text, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    p.paragraph_format.space_after = Pt(3)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


# TITLE
title = doc.add_heading("Google Ads & Meta Ads Management Proposal", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.name = "Calibri"
title.runs[0].font.size = Pt(22)
title.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
title.paragraph_format.space_after = Pt(6)

sub = doc.add_paragraph("Creekside Marketing")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.name = "Calibri"
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
sub.paragraph_format.space_after = Pt(20)

doc.add_paragraph()

# OVERVIEW
heading1("Overview")
body("Thank you for considering Creekside Marketing as your paid advertising partner. "
     "This proposal outlines our comprehensive approach to Google Ads and Meta Ads "
     "management, designed to drive measurable results for your business.")
body("Our service includes complete onboarding, strategic campaign development, and "
     "ongoing optimization to ensure your advertising investment delivers maximum return.")

# WHY CHOOSE
# NOTE 2026-05-15: per Peterson, sales-stage proposals do NOT mention internal comms
# platforms (Google Chat, Slack, etc.) — those references are for post-close materials only.
heading1("Why Choose Creekside Marketing")
bullet("3-month minimum contract. Month one is non-cancellable. A $250 cancellation fee applies if you cancel during months 2 or 3. After the initial 3 months, service converts to month-to-month with 30-day cancellation notice.",
       bold_prefix="Contract Structure:")
bullet("You communicate directly with the specialists running your campaigns, not a support inbox or account coordinator.",
       bold_prefix="Direct Access to Your Ad Managers:")
bullet("We optimize for business growth and revenue, not just vanity metrics like low cost-per-acquisition.",
       bold_prefix="Growth-Focused Optimization:")

# ONBOARDING
heading1("Onboarding Services")
body("Our comprehensive onboarding process ensures your advertising campaigns launch on a solid foundation:")

heading3("Conversion Tracking Setup")
bullet("Complete conversion tracking implementation across Google Ads and Meta Ads platforms")
bullet("CRM integration to ensure seamless data flow and accurate performance measurement")
bullet("Server-side conversion tracking implementation when required for enhanced accuracy and privacy compliance")

heading3("Comprehensive Account Audits")
bullet("Thorough review of existing Google Ads account structure, campaigns, and performance")
bullet("Complete Meta Ads account audit including campaign analysis and audience evaluation")
bullet("Identification of optimization opportunities and strategic recommendations")

heading3("90-Day Strategic Plan")
bullet("Detailed 90-day roadmap outlining campaign priorities, testing strategies, and growth initiatives")
bullet("Clear milestones and performance targets aligned with your business objectives")
bullet("Strategic framework for scaling successful campaigns and optimizing underperformers")

heading3("Creative & Technical Support")
bullet("Landing page setup and optimization for improved conversion rates")
bullet("Creative design services including ad copy, imagery, and visual assets")

# ONGOING MANAGEMENT
heading1("Ongoing Management & Optimization")
body("Once your campaigns are live, we provide continuous hands-on management with "
     "systematic optimization across all critical performance areas:")

heading3("Communication & Reporting")
bullet("Comprehensive written reports with data analysis, insights, and strategic "
       "recommendations delivered every two weeks.",
       bold_prefix="Bi-Weekly Performance Reports:")
bullet("Dedicated time to review performance, discuss results, and align on strategic "
       "priorities for the month ahead.",
       bold_prefix="Monthly Strategy Calls:")

heading3("Google Ads Optimization Activities")
body("Our team performs systematic optimization checks and adjustments on your Google Ads campaigns:")
for item in [
    ("Search Term Exclusions:", "Regular review and addition of negative keywords to prevent wasted spend on irrelevant searches"),
    ("Notification Monitoring:", "Daily review of Google Ads notifications for policy changes, account issues, or optimization recommendations"),
    ("Ad Placement Analysis:", "Review of where ads appear to exclude underperforming placements and websites"),
    ("Disapproved Ads & Policy Violations:", "Immediate resolution of any disapproved ads or policy issues to maintain campaign performance"),
    ("Budget Management:", "Regular monitoring and adjustment of campaign budgets to maximize results within your spending targets"),
    ("Conversion Tracking:", "Ongoing verification that all conversion events are firing correctly and data is accurate"),
    ("Search Keyword Performance:", "Analysis of individual keyword performance with adjustments to bids, match types, and keyword additions"),
    ("Campaign Comparisons:", "Cross-campaign analysis to identify winners and redistribute budget to highest performers"),
    ("Ad Group Optimization:", "Regular review and optimization of ad group structure, targeting, and performance"),
]:
    bullet(item[1], bold_prefix=item[0])

heading3("Meta Ads Optimization Activities")
body("Our team performs systematic optimization checks and adjustments on your Meta Ads campaigns:")
for item in [
    ("Keyword & Ads Library Research:", "Competitive analysis using Meta's Ads Library to identify successful creative approaches and messaging strategies"),
    ("Optimization Score Monitoring:", "Regular review of Meta's optimization recommendations with strategic implementation of applicable suggestions"),
    ("Ad Strength Improvement:", "Continuous enhancement of ad creative quality scores through copy and visual optimization"),
    ("Budget Management:", "Daily monitoring and strategic reallocation of campaign budgets based on performance"),
    ("Conversion Tracking:", "Ongoing verification that Pixel events are firing correctly and attribution is working properly"),
    ("Ad Schedule Optimization:", "Use of automated rules and scheduling to optimize delivery timing based on performance patterns"),
    ("Placement Analysis:", "Review of ad placements across Facebook, Instagram, Messenger, and Audience Network to exclude underperformers"),
    ("Campaign Performance Comparison:", "Cross-campaign analysis focusing on target CPA achievement and budget optimization opportunities"),
    ("Pixel Health Monitoring:", "Regular checks to ensure Meta Pixel has no errors and all conversion events are tracking properly"),
]:
    bullet(item[1], bold_prefix=item[0])

heading3("Strategic Campaign Development")
bullet("Development and launch of new campaign structures, audiences, and targeting strategies based on performance insights.",
       bold_prefix="New Campaign Launches:")
bullet("Systematic testing of ad creative, copy, audiences, bidding strategies, and landing pages to continuously improve performance.",
       bold_prefix="A/B Testing:")

# PLATFORM COVERAGE
heading1("Platform Coverage")
body("Your campaigns will reach your target audience across multiple high-value placements:")

heading3("Google Ads Placements")
for item in [
    ("Google Search:", "Capture high-intent users actively searching for your products or services"),
    ("YouTube:", "Engage audiences with video content on the world's second-largest search engine"),
    ("Google Display Network:", "Build brand awareness across millions of websites and apps"),
    ("Gmail:", "Reach potential customers in their inbox with targeted sponsored promotions"),
    ("Google Shopping:", "(For e-commerce businesses) Showcase products directly in search results with images, pricing, and merchant information"),
]:
    bullet(item[1], bold_prefix=item[0])

heading3("Meta Ads Placements")
for item in [
    ("Facebook:", "Connect with your audience across feeds, stories, and in-stream video"),
    ("Instagram:", "Engage visually-driven audiences through feed, stories, reels, and explore placements"),
    ("WhatsApp:", "Drive direct conversations and customer engagement"),
    ("Advantage+ Placements:", "Leverage Meta's automated placement optimization across Messenger, Audience Network, and other high-performing positions"),
]:
    bullet(item[1], bold_prefix=item[0])

# INVESTMENT & TERMS
doc.add_page_break()
heading1("Investment & Terms")
body("Our management fee is based on a percentage of your ad spend, designed to scale with your "
     "business. As your ad budget grows, the percentage decreases, and a monthly cap ensures your "
     "costs stay predictable even at higher spend levels.")

heading2("Pricing Structure")

ROWS = [
    ["Onboarding Fee (one-time)", "$1,500 per platform"],
    ["Monthly Management Fee",
     "$1,500 minimum per platform\n(applies until ad spend exceeds $7,500/mo)"],
    ["Variable Rate",
     "20% up to $30k spend\n15% from $30k-$60k\n10% over $60k\n(Calculated per platform)"],
    ["Monthly Cap", "$15,000 / month"],
]

col_widths = [Inches(2.6), Inches(3.6)]
tbl = doc.add_table(rows=1 + len(ROWS), cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = tbl.rows[0].cells
shade_cell(hdr_cells[0], "1A56DB")
shade_cell(hdr_cells[1], "1A56DB")
set_cell_text(hdr_cells[0], "Creekside Management Fee", bold=True, size=11,
              color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.LEFT)
hdr_cells[0].merge(hdr_cells[1])

ROW_COLORS = ["F9FAFB", "FFFFFF", "F9FAFB", "FFFFFF"]
for r_idx, row_data in enumerate(ROWS):
    row_cells = tbl.rows[r_idx + 1].cells
    for c_idx, cell_text in enumerate(row_data):
        shade_cell(row_cells[c_idx], ROW_COLORS[r_idx])
        bold = (c_idx == 0)
        set_cell_text(row_cells[c_idx], cell_text, bold=bold, size=10.5,
                      align=WD_ALIGN_PARAGRAPH.LEFT)

for row in tbl.rows:
    for i, cell in enumerate(row.cells):
        if i < len(col_widths):
            cell.width = col_widths[i]

doc.add_paragraph()

heading2("Fee Scaling")
body("The chart below shows how your monthly management fee scales with ad spend, including "
     "the tier breakpoints where the percentage decreases and the $15,000 monthly cap.")

if os.path.exists(CHART_PATH):
    doc.add_picture(CHART_PATH, width=Inches(6.3))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph("Figure 1 -- Monthly management fee by ad spend level.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    caption.runs[0].font.italic = True
    caption.paragraph_format.space_after = Pt(12)
else:
    body(f"[Chart missing -- run proposal_chart.py first to generate {CHART_PATH}]")

heading2("Payment & Contract Terms")
for item in [
    ("Ad Spend:", "You pay Google and Meta directly for all advertising costs. We never touch your ad spend budget."),
    ("Management Fees:", "Billed separately and paid directly to Creekside Marketing."),
    ("Contract Terms:", "3-month minimum contract. Month one is non-cancellable. A $250 cancellation fee applies for cancellations in months 2 or 3. After the initial 3 months, service converts to month-to-month with 30-day cancellation notice."),
    ("Account Ownership:", "All ad accounts, audiences, and creative assets remain your property. Upon cancellation, everything transfers back to you with full access."),
]:
    bullet(item[1], bold_prefix=item[0])

# NEXT STEPS
heading1("Next Steps")
body("We're excited about the opportunity to partner with you and drive meaningful results for your business.")
bullet("Review this proposal and reach out with any questions.")
bullet("Ready to move forward? Let us know which pricing plan you prefer, and we'll send over your "
       "onboarding invoice, service agreement, and onboarding instructions with everything we need to get started.")

doc.add_paragraph()
closing = doc.add_paragraph("We look forward to helping you achieve your advertising goals.")
closing.runs[0].font.name = "Calibri"
closing.runs[0].font.size = Pt(11)
closing.paragraph_format.space_before = Pt(6)

doc.add_paragraph()
sig = doc.add_paragraph("Best regards,\nThe Creekside Marketing Team")
sig.runs[0].font.name = "Calibri"
sig.runs[0].font.size = Pt(11)
sig.runs[0].font.bold = True

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
